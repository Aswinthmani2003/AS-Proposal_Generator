import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import uuid
import tempfile
import locale

# Proposal configurations
PROPOSAL_CONFIG = {
    "Make, Manychat & CRM Automation": {
        "template": "Make, Manychat & CRM Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make & Manychat Automation": {
        "template": "Make & Manychat Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Ai Calling, Make, Manychat and CRM Automation": {
        "template": "Ai Calling, Make, Manychat and CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling + CRM Integration", "AI-Price"),
            ("ManyChat & Make Automation", "MM-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling, Make & CRM Automation": {
        "template": "AI Calling, Make & CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling", "AI-Price"),
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling(Basic) & CRM Automation": {
        "template": "AI Calling(Basic) & CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling(Basic)", "AI-Price"),
            ("CRM Automation", "CC-Price")  # Changed Here
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling, Make & Manychat Automation": {
        "template": "Ai Calling, Make & Manychat Automation.docx",
        "pricing_fields": [
            ("AI Calling(Basic)", "AI-Price"),
            ("ManyChat Automation", "MC-Price"),
            ("Make Automation", "M-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Ai Calling + CRM Intergration, Make & Manychat Automation, CRM Automation": {
        "template": "Ai Calling + CRM Intergration, Make & Manychat Automation, CRM Automation.docx",
        "pricing_fields": [
            ("AI Calling + CRM Integration", "AI-Price"),
            ("ManyChat & Make Automation", "MM-Price"),
            ("CRM Automation", "CC-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "AI Calling(Basic) & CRM Automation & Email Automation": {
        "template": "AI Calling(Basic) & CRM Automation & Email Automation.docx",
        "pricing_fields": [
            ("AI Calling(Basic)", "AI-Price"),
            ("CRM Automation", "CC-Price"),
            ("Email Automation", "E-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Manychat & CRM Automation": {
        "template": "Manychat & CRM Automation.docx",
        "pricing_fields": [
            ("ManyChat Automation", "MC-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make & CRM Automation": {
        "template": "Make & CRM Automation.docx",
        "pricing_fields": [
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    },
    "Make, CRM Automation & AI Content Creation": {
        "template": "Make, CRM Automation & AI Content Creation.docx",
        "pricing_fields": [
            ("Make Automation", "M-Price"),
            ("CRM Automations", "C-Price"),
            ("AI Content Creation", "ACC-Price")
        ],
        "team_type": "general",
        "special_fields": [("VDate", "<<")]
    }
}

def apply_formatting(new_run, original_run):
    """Copy formatting from original run to new run"""
    if original_run.font.name:
        new_run.font.name = original_run.font.name
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), original_run.font.name)
    if original_run.font.size:
        new_run.font.size = original_run.font.size
    if original_run.font.color.rgb:
        new_run.font.color.rgb = original_run.font.color.rgb
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic

def replace_in_paragraph(para, placeholders):
    """Handle paragraph replacements preserving formatting"""
    original_runs = para.runs.copy()
    full_text = para.text
    for ph, value in placeholders.items():
        full_text = full_text.replace(ph, str(value))

    if full_text != para.text:
        para.clear()
        new_run = para.add_run(full_text)
        if original_runs:
            original_run = next((r for r in original_runs if r.text), None)
            if original_run:
                apply_formatting(new_run, original_run)

def replace_and_format(doc, placeholders):
    """Enhanced replacement with table cell handling"""
    # Process paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, placeholders)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    replace_in_paragraph(para, placeholders)
                else:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, placeholders)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return doc

def get_marketing_team_details():
    """Collect team composition details specifically for marketing proposals"""
    st.subheader("Marketing Team Composition")
    team_roles = {
        "Project Manager": "PM",
        "Content Writers": "CW",
        "Graphic Designer": "GD",
        "SEO Specialists": "SE",
        "Social Media Manager": "SM",
        "Ad Campaign Manager": "AC"
    }
    team_details = {}
    cols = st.columns(3)

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 3]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"marketing_team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def get_general_team_details():
    """Collect team composition for non-marketing proposals"""
    st.subheader("Team Composition")
    team_roles = {
        "Project Manager": "P1",
        "Frontend Developers": "F1",
        "Business Analyst": "B1",
        "AI/ML Developers": "A1",
        "UI/UX Members": "U1",
        "System Architect": "S1",
        "Backend Developers": "BD1",
        "AWS Developer": "AD1"
    }
    team_details = {}
    cols = st.columns(2)

    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            count = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"team_{placeholder}"
            )
            team_details[f"<<{placeholder}>>"] = str(count)
    return team_details

def remove_empty_rows(table):
    """Remove rows from the table where the second cell is empty or has no value."""
    rows_to_remove = []
    for row in table.rows:
        if len(row.cells) > 1 and row.cells[1].text.strip() == "":
            rows_to_remove.append(row)
    # Remove rows in reverse order to avoid index issues
    for row in reversed(rows_to_remove):
        table._tbl.remove(row._element)

def validate_phone_number(country, phone_number):
    """Validate phone number based on country"""
    if country.lower() == "india":
        if not phone_number.startswith("+91"):
            return False
    else:
        if not phone_number.startswith("+1"):
            return False
    return True

def format_number_with_commas(number):
    """Format number with commas (e.g., 10000 -> 10,000)"""
    return f"{number:,}"

def generate_document():
    st.title("Proposal Generator")
    base_dir = os.path.join(os.getcwd(), "templates")

    selected_proposal = st.selectbox("Select Proposal", list(PROPOSAL_CONFIG.keys()))
    config = PROPOSAL_CONFIG[selected_proposal]
    template_path = os.path.join(base_dir, config["template"])

    # Client Information
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Client Name:")
        client_email = st.text_input("Client Email:")
    with col2:
        country = st.text_input("Country:")
        client_number = st.text_input("Client Number:")
        if client_number and country:
            if not validate_phone_number(country, client_number):
                st.error(f"Phone number for {country} should start with {'+91' if country.lower() == 'india' else '+1'}")

    date_field = st.date_input("Date:", datetime.today())

    # Currency Handling
    currency = st.selectbox("Select Currency", ["USD", "INR"])
    currency_symbol = "$" if currency == "USD" else "₹"

    # Special Fields Handling
    special_data = {}
    if config.get("special_fields"):
        st.subheader("Additional Details")
        for field, wrapper in config["special_fields"]:
            if wrapper == "<<":
                placeholder = f"<<{field}>>"
                if field == "VDate":
                    vdate = st.date_input("Proposal Validity Until:")
                    special_data[placeholder] = vdate.strftime("%d-%m-%Y")
                else:
                    special_data[placeholder] = st.text_input(f"{field.replace('_', ' ').title()}:")

    # Pricing Section
    st.subheader("Pricing Details")
    pricing_data = {}
    numerical_values = {}  # To store raw numerical values for calculations

    # Determine number of columns based on selected proposal
    num_pricing_fields = len(config["pricing_fields"])
    cols = st.columns(num_pricing_fields)

    # Collect base services input
    for idx, (label, key) in enumerate(config["pricing_fields"]):
        with cols[idx % num_pricing_fields]:
            value = st.number_input(
                f"{label} ({currency})",
                min_value=0,
                value=0,
                step=100,
                format="%d",
                key=f"price_{key}"
            )
            numerical_values[key] = value
            # Only add to pricing_data if the value is greater than 0
            if value > 0:
                pricing_data[f"<<{key}>>"] = f"{currency_symbol}{format_number_with_commas(value)}"
            else:
                pricing_data[f"<<{key}>>"] = ""  # Empty value for fields with zero values

    # Calculate Total Amount
    if selected_proposal == "Make, Manychat & CRM Automation":
        # Calculate sum of services
        services_sum = sum([
            numerical_values.get("MC-Price", 0),
            numerical_values.get("M-Price", 0),
            numerical_values.get("C-Price", 0)
        ])
    elif selected_proposal == "Make & Manychat Automation":
        # Calculate sum of services
        services_sum = sum([
            numerical_values.get("MC-Price", 0),
            numerical_values.get("M-Price", 0)
        ])
    elif selected_proposal == "Ai Calling, Make, Manychat and CRM Automation":
        services_sum = sum([
            numerical_values.get("AI-Price", 0),
            numerical_values.get("MM-Price", 0)
        ])
    elif selected_proposal == "AI Calling, Make & CRM Automation":
        services_sum = sum([
            numerical_values.get("AI-Price", 0),
            numerical_values.get("M-Price", 0),
            numerical_values.get("C-Price", 0)
        ])
    elif selected_proposal == "AI Calling(Basic) & CRM Automation":
        services_sum = sum([
            numerical_values.get("AI-Price", 0),
            numerical_values.get("CC-Price", 0) # Changed Here
        ])
    elif selected_proposal == "AI Calling, Make & Manychat Automation":
        services_sum = sum([
            numerical_values.get("AI-Price", 0),
            numerical_values.get("MC-Price", 0),
            numerical_values.get("M-Price", 0)
        ])
    elif selected_proposal == "Ai Calling + CRM Intergration, Make & Manychat Automation, CRM Automation":
        services_sum = sum([
            numerical_values.get("AI-Price", 0),
            numerical_values.get("MM-Price", 0),
            numerical_values.get("CC-Price", 0)
        ])
    elif selected_proposal == "AI Calling(Basic) & CRM Automation & Email Automation":
        services_sum = sum([
            numerical_values.get("AI-Price", 0),
            numerical_values.get("CC-Price", 0),
            numerical_values.get("E-Price", 0)
        ])
    elif selected_proposal == "Manychat & CRM Automation":
        services_sum = sum([
            numerical_values.get("MC-Price", 0),
            numerical_values.get("C-Price", 0)
        ])
    elif selected_proposal == "Make & CRM Automation":
        services_sum = sum([
            numerical_values.get("M-Price", 0),
            numerical_values.get("C-Price", 0)
        ])
    elif selected_proposal == "Make, CRM Automation & AI Content Creation":
        services_sum = sum([
            numerical_values.get("M-Price", 0),
            numerical_values.get("C-Price", 0),
            numerical_values.get("ACC-Price", 0)
        ])

    # Annual Maintenance (10% of Total Amount)
    am_price = int(services_sum * 0.10)
    pricing_data["<<AM-Price>>"] = f"{currency_symbol}{format_number_with_commas(am_price)}"

    # Total Amount
    total = services_sum + am_price
    if currency == "INR":
        pricing_data["<<T-Price>>"] = f"{currency_symbol}{format_number_with_commas(total)} + 18% GST"
    else:
        pricing_data["<<T-Price>>"] = f"{currency_symbol}{format_number_with_commas(total)}"

    # Additional Features & Enhancements
    af_price = 250 if currency == "USD" else 25000
    pricing_data["<<AF-Price>>"] = f"{currency_symbol}{format_number_with_commas(af_price)}"

    # Team Composition
    team_data = {}
    if config["team_type"] == "marketing":
        team_data = get_marketing_team_details()
    elif config["team_type"] == "general":
        team_data = get_general_team_details()

    # Add Additional Tools Section
    st.subheader("Add Additional Tools")
    additional_tool_1 = st.text_input("Tool 1:")
    additional_tool_2 = st.text_input("Tool 2:")

    additional_tools_data = {}
    if additional_tool_1:
        additional_tools_data["<<T1>>"] = additional_tool_1
    else:
        additional_tools_data["<<T1>>"] = ""  # Make placeholder invisible

    if additional_tool_2:
        additional_tools_data["<<T2>>"] = additional_tool_2
    else:
        additional_tools_data["<<T2>>"] = ""  # Make placeholder invisible
   

    # Combine all placeholders
    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Email>>": client_email,
        "<<Client Number>>": client_number,
        "<<Date>>": date_field.strftime("%d-%m-%Y"),
        "<<Country>>": country
    }
    placeholders.update(pricing_data)
    placeholders.update(team_data)
    placeholders.update(special_data)
    placeholders.update(additional_tools_data)

    if st.button("Generate Proposal"):
        if client_number and country and not validate_phone_number(country, client_number):
            st.error(f"Invalid phone number format for {country} should start with {'+91' if country.lower() == 'india' else '+1'}.")
        else:
            formatted_date = date_field.strftime("%d %b %Y")
            unique_id = str(uuid.uuid4())[:8]
            doc_filename = f"{selected_proposal}_{client_name}_{formatted_date}_{unique_id}.docx"

            with tempfile.TemporaryDirectory() as temp_dir:
                doc = Document(template_path)
                doc = replace_and_format(doc, placeholders)

                # Remove empty rows from the pricing table
                for table in doc.tables:
                    remove_empty_rows(table)

                doc_path = os.path.join(temp_dir, doc_filename)
                doc.save(doc_path)

                with open(doc_path, "rb") as f:
                    st.download_button(
                        label="Download Proposal",
                        data=f,
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    generate_document()
